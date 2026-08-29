from pathlib import Path

from lesson_plan_ai.models.evidence import CurriculumRecord


def extract_scheme(path: str | Path) -> list[CurriculumRecord]:
    """Extract paragraphs and tables while retaining a human-readable location."""
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required to parse scheme documents") from exc
    document = Document(path)
    records: list[CurriculumRecord] = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            records.append(CurriculumRecord(
                source_document=str(path), source_location=f"paragraph {index}", text=text,
                unit_topic=paragraph.style.name if paragraph.style else None,
            ))
    for table_index, table in enumerate(document.tables, start=1):
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        text = "\n".join(row for row in rows if row)
        if text:
            records.append(CurriculumRecord(
                source_document=str(path), source_location=f"table {table_index}", text=text,
            ))
    return records
