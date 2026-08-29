from io import BytesIO

from lesson_plan_ai.models.lesson_plan import LessonPlan


def generate_docx(plan: LessonPlan, template_path: str | None = None) -> bytes:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required to generate lesson documents") from exc
    document = Document(template_path) if template_path else Document()
    document.add_heading(plan.title, level=1)
    document.add_paragraph(f"{plan.subject} | {plan.year_group} | {plan.duration_minutes} minutes")
    document.add_heading("Learning objectives", level=2)
    for objective in plan.learning_objectives:
        document.add_paragraph(objective.objective, style="List Bullet")
    document.add_heading("Lesson activities", level=2)
    table = document.add_table(rows=1, cols=4)
    for cell, heading in zip(table.rows[0].cells, ("Activity", "Minutes", "Teacher", "Students")):
        cell.text = heading
    for activity in plan.activities:
        cells = table.add_row().cells
        cells[0].text, cells[1].text = activity.name, str(activity.minutes)
        cells[2].text, cells[3].text = activity.teacher_actions, activity.student_actions
    document.add_heading("Differentiation", level=2)
    document.add_paragraph("Support: " + "; ".join(plan.differentiation.support))
    document.add_paragraph("Stretch: " + "; ".join(plan.differentiation.stretch))
    document.add_heading("Assessment", level=2)
    document.add_paragraph(f"{plan.assessment.method}: {plan.assessment.expected_evidence}")
    if plan.sources:
        document.add_heading("Sources", level=2)
        for source in plan.sources:
            document.add_paragraph(f"{source.title}: {source.provenance}", style="List Bullet")
    output = BytesIO()
    document.save(output)
    return output.getvalue()
